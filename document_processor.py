# document_processor.py

import io
import fitz  # PyMuPDF for PDF processing
import pdfplumber
from PIL import Image
from docx import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from pathlib import Path
from config import IMAGES_DIR  # Use centralized config

# Ensure the images directory exists
IMAGES_DIR.mkdir(parents=True, exist_ok=True)


class DocumentProcessor:
    def __init__(self, images_dir=IMAGES_DIR):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            separators=["\n\n", "\n", " ", ""]
        )
        self.images_dir = images_dir

    def extract_text_images_tables_pdf(self, file_path: str):
        """Extract text, images, and tables from a PDF file."""
        text = ""
        image_refs = []
        tables = []

        # Extract text and images using PyMuPDF
        doc = fitz.open(file_path)
        for page_number in range(len(doc)):
            page = doc[page_number]
            text += page.get_text()

            for img_index, img in enumerate(page.get_images(full=True)):
                xref = img[0]
                base_image = doc.extract_image(xref)
                image_bytes = base_image['image']
                ext = base_image['ext']
                image_obj = Image.open(io.BytesIO(image_bytes))
                image_name = f"pdf_page{page_number + 1}_img{img_index + 1}.{ext}"
                image_path = self.images_dir / image_name
                image_obj.save(image_path)
                image_refs.append({
                    "type": "image",
                    "file": str(image_path),
                    "caption": f"Image on page {page_number + 1}",
                    "page": page_number + 1
                })
        doc.close()

        # Extract tables using pdfplumber
        with pdfplumber.open(file_path) as pdf:
            for page_number, page in enumerate(pdf.pages):
                for table_index, table in enumerate(page.extract_tables()):
                    if table:
                        tables.append({
                            "type": "table",
                            "content": table,
                            "caption": f"Table {table_index + 1} on page {page_number + 1}",
                            "page": page_number + 1
                        })

        return text, image_refs, tables

    def extract_text_images_tables_docx(self, file_path: str):
        """Extract text, images, and tables from a Word (.docx) file."""
        doc = Document(file_path)
        text = ""
        tables_list = []
        image_refs = []

        # Extract paragraphs
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"

        # Extract tables
        for table_index, table in enumerate(doc.tables):
            table_content = [[cell.text for cell in row.cells] for row in table.rows]
            tables_list.append({
                "type": "table",
                "content": table_content,
                "caption": f"Table {table_index + 1} in document",
                "table_index": table_index
            })

        # Extract images (from document relationships)
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                img_bytes = rel.target_part.blob
                img_ext = rel.target_ref.split('.')[-1]
                img_name = f"{Path(file_path).stem}_{rel.target_ref.split('/')[-1]}"
                img_path = self.images_dir / img_name
                with open(img_path, 'wb') as f:
                    f.write(img_bytes)
                image_refs.append({
                    "type": "image",
                    "file": str(img_path),
                    "caption": f"Image in document: {img_name}"
                })

        return text, image_refs, tables_list

    def process_document(self, file_path: str, file_type: str):
        """
        Process the input document (PDF/DOCX), split into chunks, and attach metadata
        for images/tables found within.

        Returns:
            A list of dicts with keys: "type", "content", and "metadata".
        """
        chunks_with_meta = []

        if file_type.lower() == 'pdf':
            text, images, tables = self.extract_text_images_tables_pdf(file_path)
        elif file_type.lower() in ['docx', 'doc']:
            text, images, tables = self.extract_text_images_tables_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")

        # Append image chunks
        for image in images:
            chunks_with_meta.append({
                "type": "image",
                "content": "",
                "metadata": image
            })

        # Append table chunks
        for table in tables:
            chunks_with_meta.append({
                "type": "table",
                "content": table["content"],
                "metadata": table
            })

        # Split and append text chunks
        text_chunks = self.text_splitter.split_text(text)
        for idx, chunk in enumerate(text_chunks):
            chunks_with_meta.append({
                "type": "text",
                "content": chunk,
                "metadata": {"chunk_id": idx}
            })

        return chunks_with_meta
